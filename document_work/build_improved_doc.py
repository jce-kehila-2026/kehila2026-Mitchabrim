from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "المستند_الختامي_للجمعية_מתחברים_محسن.docx"

BURGUNDY = "7A2333"
GOLD = "C69A4B"
INK = "24313A"
MUTED = "5F6B72"
PALE = "F7F1E8"
PALE_BLUE = "EEF4F5"
WHITE = "FFFFFF"


def set_cell_or_paragraph_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, name="Arial", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().set(qn("w:rtl"), "1")


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side="right", color=BURGUNDY, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    tag = qn(f"w:{side}")
    border = p_bdr.find(tag)
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url, color=BURGUNDY):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:cs"), "Arial")
    run_pr.append(r_fonts)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    run_pr.append(rtl)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_pr.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(underline)
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("صفحة ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, BURGUNDY, 16, 8),
        ("Heading 2", 13, INK, 11, 5),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "FAQ Question" not in styles:
        faq = styles.add_style("FAQ Question", WD_STYLE_TYPE.PARAGRAPH)
    else:
        faq = styles["FAQ Question"]
    faq.font.name = "Arial"
    faq.font.size = Pt(11.5)
    faq.font.bold = True
    faq.font.color.rgb = RGBColor.from_string(BURGUNDY)
    faq.paragraph_format.space_before = Pt(7)
    faq.paragraph_format.space_after = Pt(4)
    faq.paragraph_format.keep_with_next = True

    if "Lead Callout" not in styles:
        callout = styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Lead Callout"]
    callout.font.name = "Arial"
    callout.font.size = Pt(11)
    callout.font.bold = True
    callout.font.color.rgb = RGBColor.from_string(INK)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    set_cell_or_paragraph_rtl(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True, color=BURGUNDY)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=INK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_cell_or_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BURGUNDY if level == 1 else INK)
    if level == 1:
        set_paragraph_border(paragraph, side="bottom", color=GOLD, size="8", space="4")
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_cell_or_paragraph_rtl(paragraph)
    paragraph.paragraph_format.right_indent = Inches(0.23)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, color=INK)
    return paragraph


def add_callout(doc, label, text, fill=PALE_BLUE):
    paragraph = doc.add_paragraph(style="Lead Callout")
    set_cell_or_paragraph_rtl(paragraph)
    set_paragraph_shading(paragraph, fill)
    set_paragraph_border(paragraph, side="right", color=BURGUNDY, size="20", space="8")
    first = paragraph.add_run(f"{label}: ")
    set_run_font(first, bold=True, color=BURGUNDY)
    second = paragraph.add_run(text)
    set_run_font(second, color=INK)
    return paragraph


def add_faq(doc, question, answer):
    q = doc.add_paragraph(style="FAQ Question")
    set_cell_or_paragraph_rtl(q)
    set_paragraph_shading(q, PALE)
    set_paragraph_border(q, side="right", color=GOLD, size="18", space="8")
    run = q.add_run(question)
    set_run_font(run, size=11.5, bold=True, color=BURGUNDY)

    a = doc.add_paragraph()
    set_cell_or_paragraph_rtl(a)
    a.paragraph_format.right_indent = Inches(0.18)
    a.paragraph_format.space_after = Pt(7)
    label = a.add_run("الإجابة: ")
    set_run_font(label, bold=True, color=BURGUNDY)
    text = a.add_run(answer)
    set_run_font(text, color=INK)


def add_contact(doc, name, email):
    paragraph = doc.add_paragraph()
    set_cell_or_paragraph_rtl(paragraph)
    set_paragraph_shading(paragraph, PALE_BLUE)
    set_paragraph_border(paragraph, side="right", color=BURGUNDY, size="18", space="8")
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    name_run = paragraph.add_run(f"{name}\n")
    set_run_font(name_run, size=11.5, bold=True, color=BURGUNDY)
    add_hyperlink(paragraph, email, f"mailto:{email}")


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    configure_styles(doc)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    set_cell_or_paragraph_rtl(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("מתחברים  |  دليل التسليم والاستخدام")
    set_run_font(hr, size=8.5, bold=True, color=MUTED)
    set_paragraph_border(hp, side="bottom", color=GOLD, size="6", space="4")

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)

    # Cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_or_paragraph_rtl(title)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    tr = title.add_run("المستند الختامي للجمعية")
    set_run_font(tr, size=27, bold=True, color=BURGUNDY)

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(7)
    pr = project.add_run("مشروع מתחברים")
    set_run_font(pr, size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    sr = subtitle.add_run("منصة لإدارة المتطوعين ودعم كبار السن")
    set_run_font(sr, size=13, color=MUTED)

    accent = doc.add_paragraph()
    accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    accent.paragraph_format.left_indent = Inches(1.7)
    accent.paragraph_format.right_indent = Inches(1.7)
    set_paragraph_border(accent, side="bottom", color=GOLD, size="18", space="3")

    link_label = doc.add_paragraph()
    link_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_label.paragraph_format.space_before = Pt(22)
    lr = link_label.add_run("رابط الموقع الرسمي")
    set_run_font(lr, size=10.5, bold=True, color=MUTED)
    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(link_p, "mitchabrim-jce2026.web.app", "https://mitchabrim-jce2026.web.app")

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.paragraph_format.space_before = Pt(34)
    rr = prepared.add_run("إعداد فريق المشروع  |  2026")
    set_run_font(rr, size=10.5, color=MUTED)

    doc.add_page_break()

    add_heading(doc, "محتويات المستند", 1)
    for item in [
        "1. مقدمة عن النظام",
        "2. الجمهور المستهدف",
        "3. متطلبات النظام",
        "4. الوصول إلى النظام",
        "5. تعليمات استخدام المسؤول",
        "6. تعليمات استخدام المتطوع",
        "7. إدارة الصلاحيات والوصول",
        "8. الأسئلة الشائعة وحل المشكلات",
        "9. الدعم بعد التسليم",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "ملاحظة",
        "هذا المستند دليل تشغيلي مختصر. يُنصح بتحديثه عند إضافة خصائص أو تغيير الصلاحيات أو إجراءات العمل.",
    )

    add_heading(doc, "1. مقدمة عن النظام", 1)
    add_body(
        doc,
        "מתחברים هو نظام ويب متكامل طُوّر لمساعدة الجمعية على إدارة عملها مع كبار السن والمتطوعين بصورة منظمة من خلال منصة واحدة. يجمع النظام بين موقع عام للتعريف بالجمعية، ولوحة إدارة للمسؤولين، ومنطقة شخصية للمتطوعين.",
    )
    add_body(
        doc,
        "يهدف النظام إلى تقليل الاعتماد على الملفات المتفرقة والإجراءات اليدوية، وتسهيل متابعة البيانات والمهام والتقارير والمشاريع ومحتوى الموقع العام.",
    )

    add_heading(doc, "2. الجمهور المستهدف", 1)
    add_body(doc, "يخدم النظام ثلاث فئات رئيسية:")
    add_bullet(doc, "مسؤولو الجمعية: لإدارة كبار السن والمتطوعين والمشاريع والتقارير والمهام والمحتوى والصلاحيات.")
    add_bullet(doc, "المتطوعون: لمتابعة المهام، وإرسال تقارير اللقاءات أو الاتصالات، ومراجعة البيانات الشخصية.")
    add_bullet(doc, "زوار الموقع العام: للتعرّف إلى الجمعية ونشاطاتها وشركائها وإرسال طلبات الانضمام.")

    add_heading(doc, "3. متطلبات النظام", 1)
    add_body(
        doc,
        "يعمل النظام عبر المتصفح ولا يحتاج إلى تثبيت برنامج خاص. يلزم اتصال بالإنترنت ومتصفح حديث مثل Google Chrome أو Microsoft Edge أو Firefox أو Safari.",
    )
    add_body(
        doc,
        "واجهة النظام الأساسية باللغة العبرية وتعمل من اليمين إلى اليسار. الموقع العام متاح للجميع، بينما تتطلب لوحة الإدارة ومنطقة المتطوع حسابًا فعّالًا وصلاحية مناسبة.",
    )
    add_callout(
        doc,
        "توصية",
        "تُفضّل شاشة حاسوب أو جهاز لوحي للمهام الإدارية التي تتضمن جداول كبيرة، مع إمكانية استخدام الهاتف للمهام السريعة.",
        fill=PALE,
    )

    add_heading(doc, "4. الوصول إلى النظام", 1)
    add_body(doc, "يمكن الوصول إلى النسخة المنشورة من خلال الرابط التالي:")
    p = doc.add_paragraph()
    set_cell_or_paragraph_rtl(p)
    add_hyperlink(p, "فتح الموقع الرسمي", "https://mitchabrim-jce2026.web.app")
    add_body(
        doc,
        "يدخل الزائر إلى الموقع العام مباشرة. أما المسؤول أو المتطوع فيستخدم صفحة تسجيل الدخول بالحساب الذي أُنشئ له مسبقًا. يجب أن يمتلك كل مستخدم حسابًا مستقلًا، ولا يجوز مشاركة حسابات الإدارة أو كلمات المرور.",
    )

    add_heading(doc, "5. تعليمات استخدام المسؤول", 1)
    add_body(
        doc,
        "بعد تسجيل الدخول بحساب إداري، تظهر لوحة الإدارة التي تجمع الأقسام الأساسية للنظام. فيما يلي وظيفة كل قسم وما يحتاج المسؤول إلى متابعته:",
    )

    admin_sections = [
        (
            "5.1 إدارة كبار السن",
            "يُستخدم هذا القسم لإضافة السجلات وتعديلها والبحث فيها وتصفيتها وأرشفتها. يمكن ربط الشخص بمنطقة وحي ومتطوع، وتحديث بيانات الاتصال والحالة والملاحظات عند الحاجة.",
        ),
        (
            "5.2 إدارة المتطوعين والمجموعات",
            "يتيح هذا القسم إضافة المتطوعين وتحديث بياناتهم وحالتهم، وتنظيم بعضهم ضمن مجموعات، ومراجعة ارتباطاتهم بكبار السن والمشاريع.",
        ),
        (
            "5.3 المشاريع والتوزيعات",
            "يُستخدم لإنشاء المشاريع وتحديد التاريخ والسنة، وإضافة المشاركين من كبار السن أو مجموعات المتطوعين، ومتابعة التنفيذ، وعرض القوائم أو طباعتها عند الحاجة.",
        ),
        (
            "5.4 البرلمانات والمعاملات المالية",
            "يتيح إدارة البرلمانات ومواقعها ومنسقيها وأعضائها ومواعيد اللقاءات والحضور والملاحظات، إضافة إلى متابعة المعاملات المالية المرتبطة بكل برلمان.",
        ),
        (
            "5.5 تقارير المتطوعين والمهام",
            "يمكن للمسؤول مراجعة تقارير المتطوعين بحسب المتطوع والتاريخ، والاطلاع على تفاصيل اللقاءات، ومتابعة الحالات التي تحتاج إلى إجراء إضافي. كما يمكن إسناد مهام لمتطوع واحد أو لعدة متطوعين ومتابعة حالتها.",
        ),
        (
            "5.6 المؤسسات وجهات الاتصال",
            "يُستخدم لإدارة المؤسسات والأشخاص المرتبطين بها، وإدارة جهات اتصال كبار السن مثل الجار أو أحد أفراد العائلة أو العامل الاجتماعي.",
        ),
        (
            "5.7 طلبات الانضمام",
            "تظهر الطلبات التي أرسلها الزوار من الموقع العام، ويستطيع المسؤول مراجعتها واتخاذ الإجراء المناسب وفق سياسة الجمعية.",
        ),
        (
            "5.8 إدارة محتوى الموقع العام",
            "يتيح تحديث الشركاء ونشاطات الجمعية وأعضاء الفريق والصور والمعرض والمحتوى الظاهر للزوار دون الحاجة إلى تعديل الكود.",
        ),
        (
            "5.9 إعدادات النظام",
            "يتيح إدارة المناطق والأحياء وفئات الصور والروابط، ومنح صلاحيات الوصول للمسؤولين والمتطوعين. يجب تنفيذ تغييرات المناطق والأحياء بحذر لأنها قد تكون مرتبطة بسجلات قائمة.",
        ),
    ]
    for heading, text in admin_sections:
        add_heading(doc, heading, 2)
        add_body(doc, text)

    add_heading(doc, "6. تعليمات استخدام المتطوع", 1)
    add_body(doc, "بعد تسجيل الدخول، يصل المتطوع إلى منطقته الشخصية، ومنها يستطيع:")
    for item in [
        "مشاهدة ملخص الحساب وعدد المهام المفتوحة.",
        "فتح المهام المرسلة إليه وقراءة التفاصيل المطلوبة.",
        "إرسال تقرير بعد لقاء أو اتصال، مع تحديد الشخص، ونوع اللقاء، والتاريخ، ووصف مختصر.",
        "مراجعة التقارير السابقة التي أرسلها.",
        "مراجعة البيانات الشخصية وإرسال طلب تحديث عند الحاجة.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. إدارة الصلاحيات والوصول", 1)
    add_body(doc, "يوجد مستويان من الصلاحيات: صلاحيات داخل التطبيق، وصلاحيات على الموارد التقنية.")
    add_bullet(doc, "داخل التطبيق: تُحدد الصلاحيات بحسب نوع المستخدم وحالة الحساب.")
    add_bullet(doc, "في Firebase: تُدار صلاحيات قاعدة البيانات والمصادقة والتخزين والاستضافة والوظائف السحابية.")
    add_bullet(doc, "في GitHub: تُدار صلاحيات الوصول إلى الكود والمستودع وعمليات التطوير.")
    add_body(
        doc,
        "يجب مراجعة الصلاحيات دوريًا، وتعطيل وصول أي شخص غادر الجمعية أو لم يعد بحاجة إليه، وتفعيل التحقق بخطوتين للحسابات الإدارية متى أمكن.",
    )
    add_callout(
        doc,
        "حماية البيانات",
        "تُستخدم الحسابات الشخصية فقط، ولا تُرسل بيانات كبار السن عبر قنوات غير معتمدة، ولا تُشارك كلمات المرور أو رموز التحقق.",
        fill=PALE,
    )

    add_heading(doc, "8. الأسئلة الشائعة وحل المشكلات", 1)
    faq_items = [
        (
            "لا أستطيع تسجيل الدخول. ماذا أفعل؟",
            "تحقق من كتابة البريد الإلكتروني وكلمة المرور بصورة صحيحة. إذا استمرت المشكلة، اطلب من المسؤول التأكد من أن الحساب موجود وحالته «فعّال» وأن دوره صحيح.",
        ),
        (
            "نسيت كلمة المرور. كيف أستعيدها؟",
            "استخدم رابط «نسيت كلمة المرور» في صفحة الدخول، ثم افتح رسالة الاستعادة المرسلة إلى بريدك. إذا لم تصل الرسالة، افحص مجلد الرسائل غير المرغوب فيها وتأكد من صحة البريد المسجل.",
        ),
        (
            "المتطوع لا يرى أحد كبار السن المرتبطين به.",
            "تأكد من ربط سجل الشخص بالمتطوع الصحيح، ومن أن حساب الدخول مرتبط بسجل المتطوع نفسه، وأن الحساب فعّال.",
        ),
        (
            "تظهر رسالة تفيد بعدم وجود صلاحية.",
            "سجّل الخروج ثم ادخل مجددًا. إذا استمرت الرسالة، يجب على مسؤول النظام التحقق من دور الحساب وحالته؛ لا تحاول استخدام حساب شخص آخر.",
        ),
        (
            "التعديل لا يظهر بعد الحفظ.",
            "انتظر رسالة نجاح الحفظ، ثم حدّث الصفحة مرة واحدة. إذا لم يظهر التعديل، لا تكرر العملية عدة مرات؛ سجّل وقت المشكلة والقسم المستخدم وأرسل التفاصيل إلى الدعم.",
        ),
        (
            "الصورة لا تظهر في الموقع العام.",
            "تأكد من اكتمال رفع الصورة ومن تفعيل خيار ظهورها في الموقع العام. إذا ظهرت رسالة نجاح ولم تظهر الصورة بعد التحديث، أرسل اسم الصورة ووقت الرفع إلى الدعم.",
        ),
        (
            "فشل إرسال تقرير أو طلب.",
            "تحقق من الاتصال بالإنترنت ومن تعبئة الحقول المطلوبة، ثم أعد المحاولة مرة واحدة. إذا تكرر الفشل، احتفظ بنص التقرير مؤقتًا وأرسل وقت الخطأ ونوع العملية إلى المسؤول.",
        ),
        (
            "هل يمكن استخدام النظام من الهاتف؟",
            "نعم، يعمل النظام في متصفح حديث على الهاتف. لكن يفضّل استخدام حاسوب أو جهاز لوحي للعمليات الإدارية التي تتضمن جداول كبيرة أو طباعة.",
        ),
        (
            "لماذا لا يمكن حذف حي أو منطقة؟",
            "يمنع النظام الحذف إذا كانت هناك سجلات مرتبطة، حتى لا تصبح البيانات غير متناسقة. يجب أولًا نقل السجلات أو تحديث المنطقة والحي، ثم إعادة محاولة الحذف.",
        ),
        (
            "من يستطيع الاطلاع على بيانات كبار السن؟",
            "المسؤولون الفعّالون، والمتطوع الذي تسمح له صلاحياته برؤية السجل المرتبط به فقط. يجب عدم مشاركة البيانات خارج النظام إلا وفق سياسة الجمعية.",
        ),
        (
            "كيف نمنع موظفًا أو متطوعًا سابقًا من الدخول؟",
            "عطّل حسابه فورًا داخل النظام، ثم اطلب من المسؤول التقني إزالة وصوله من أنظمة الإدارة والكود، وراجع الحسابات المرتبطة والمهام المفتوحة.",
        ),
    ]
    for question, answer in faq_items:
        add_faq(doc, question, answer)

    add_heading(doc, "9. الدعم بعد التسليم", 1)
    add_body(
        doc,
        "يمكن للجمعية التواصل مع أعضاء فريق التطوير للحصول على دعم تقني متعلق بتشغيل النظام ومعالجة المشكلات الأساسية.",
    )
    add_heading(doc, "جهات التواصل", 2)
    add_contact(doc, "بهاء عقل", "gonatsu123@gmail.com")
    add_contact(doc, "أحمد بكري", "ahmadbk179@gmail.com")
    add_callout(
        doc,
        "نطاق الدعم",
        "يُقدَّم الدعم الفني حسب الحاجة والإمكانية، ولا يُعد التزامًا دائمًا أو عقد صيانة مستمرًا. ويمكن الاتفاق مستقبلًا بصورة منفصلة على أي تطوير إضافي أو صيانة مستمرة، مع تحديد نطاق العمل والمدة والمسؤوليات والتكاليف.",
        fill=PALE,
    )

    # Normalize direct formatting and RTL on all text-bearing paragraphs.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() and paragraph.alignment not in {
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
        }:
            set_cell_or_paragraph_rtl(paragraph)
        for run in paragraph.runs:
            if run.text:
                if run.font.name is None:
                    set_run_font(run)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
