from pathlib import Path
from docx import Document

path = Path(__file__).with_name("original.docx")
doc = Document(path)

print(f"sections={len(doc.sections)} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
for index, section in enumerate(doc.sections):
    print(
        "SECTION",
        index,
        f"size={section.page_width}x{section.page_height}",
        f"margins={section.top_margin},{section.right_margin},{section.bottom_margin},{section.left_margin}",
    )

for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{index:03d} [{paragraph.style.name}] align={paragraph.alignment}: {text}")

for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)} style={table.style.name if table.style else None}")
    for row_index, row in enumerate(table.rows):
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        print(f"  R{row_index:02d}: {cells}")
